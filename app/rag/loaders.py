import os
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader


@dataclass
class Section:
    heading: str
    text: str
    loc: str


@dataclass
class LoadedDocument:
    doc_id: str
    source_path: str
    title: str
    sections: List[Section]
    version: str = "unknown"


def _generate_doc_id(file_path: Path) -> str:
    return f"{file_path.stem}"


def load_pdf(path: Path) -> LoadedDocument:
    reader = PdfReader(path)
    sections: List[Section] = []
    for idx, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        sections.append(
            Section(
                heading=f"Página {idx}",
                text=text.strip(),
                loc=f"page:{idx}",
            )
        )

    return LoadedDocument(
        doc_id=_generate_doc_id(path),
        source_path=str(path),
        title=path.stem,
        sections=sections,
    )


def load_docx(path: Path) -> LoadedDocument:
    doc = DocxDocument(path)
    sections: List[Section] = []
    current_heading = "Intro"
    buffer: List[str] = []

    def flush(idx: int) -> None:
        if buffer:
            sections.append(
                Section(
                    heading=current_heading,
                    text=" ".join(buffer).strip(),
                    loc=f"paragraph:{idx}",
                )
            )

    paragraph_idx = 0
    for block in doc.paragraphs:
        paragraph_idx += 1
        style = block.style.name.lower()
        text = block.text.strip()
        if not text:
            continue
        if "heading" in style:
            flush(paragraph_idx - 1)
            buffer = []
            current_heading = text[:60]
            continue
        buffer.append(text)

    flush(paragraph_idx)

    return LoadedDocument(
        doc_id=_generate_doc_id(path),
        source_path=str(path),
        title=path.stem,
        sections=sections,
    )


def load_html(path: Path) -> LoadedDocument:
    with path.open(encoding="utf-8") as handle:
        soup = BeautifulSoup(handle, "html.parser")

    sections: List[Section] = []
    for idx, element in enumerate(
        soup.select("h1, h2, h3, h4, h5, h6, p, li"), start=1
    ):
        heading = element.get_text(strip=True)
        if not heading:
            continue
        parent = element
        if element.name in {"li"}:
            parent = element.find_parent("ul") or element
        text = parent.get_text(separator=" ", strip=True)
        sections.append(
            Section(
                heading=heading,
                text=text,
                loc=f"html:{idx}",
            )
        )

    title = soup.title.string if soup.title else path.stem
    return LoadedDocument(
        doc_id=_generate_doc_id(path),
        source_path=str(path),
        title=title,
        sections=sections,
    )


def discover_documents(root: Path) -> Iterable[LoadedDocument]:
    for entry in root.iterdir():
        if entry.is_dir():
            yield from discover_documents(entry)
            continue

        if entry.suffix.lower() == ".pdf":
            yield load_pdf(entry)
        elif entry.suffix.lower() == ".docx":
            yield load_docx(entry)
        elif entry.suffix.lower() == ".html":
            yield load_html(entry)
        else:
            continue
